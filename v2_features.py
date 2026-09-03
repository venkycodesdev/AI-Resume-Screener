"""Version 2 features for the AI Resume Screener.

The module is registered from ``app.py`` after the Version 1 models and scoring
helpers exist. New data is stored in separate tables, which keeps existing
Version 1 databases compatible with ``db.create_all()``.
"""

import json
import re
from datetime import datetime, timezone
from io import BytesIO
from types import SimpleNamespace

from docx import Document
from flask import Blueprint, flash, redirect, render_template, request, send_file, url_for
from flask_login import current_user, login_required
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from werkzeug.utils import secure_filename


DEFAULT_WEIGHTS = {
    "skills": 40,
    "experience": 30,
    "projects": 20,
    "education": 10,
}
ALLOWED_TONES = {"professional", "concise", "technical", "fresher"}


def _json(value, fallback):
    try:
        parsed = json.loads(value or "")
        return parsed if isinstance(parsed, type(fallback)) else fallback
    except (TypeError, json.JSONDecodeError):
        return fallback


def _clean_weights(value):
    if isinstance(value, str):
        value = _json(value, {})
    if not isinstance(value, dict):
        return DEFAULT_WEIGHTS.copy()
    try:
        weights = {key: int(value.get(key, 0)) for key in DEFAULT_WEIGHTS}
    except (TypeError, ValueError):
        return DEFAULT_WEIGHTS.copy()
    if any(number < 0 or number > 100 for number in weights.values()):
        return DEFAULT_WEIGHTS.copy()
    if sum(weights.values()) != 100:
        return DEFAULT_WEIGHTS.copy()
    return weights


def register_v2_features(app, db, User, Analysis, helpers):
    """Register V2 models, routes and services with the existing Flask app."""

    class ScoringPreference(db.Model):
        __tablename__ = "scoring_preference"
        id = db.Column(db.Integer, primary_key=True)
        user_id = db.Column(db.Integer, db.ForeignKey("user.id"), unique=True, nullable=False, index=True)
        weights = db.Column(db.Text, nullable=False, default=lambda: json.dumps(DEFAULT_WEIGHTS))
        updated_at = db.Column(db.DateTime, nullable=False, default=lambda: datetime.now(timezone.utc))

    class AnalysisContext(db.Model):
        __tablename__ = "analysis_context"
        id = db.Column(db.Integer, primary_key=True)
        user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False, index=True)
        analysis_id = db.Column(db.Integer, db.ForeignKey("analysis.id", ondelete="CASCADE"), unique=True, nullable=False, index=True)
        resume_text = db.Column(db.Text, nullable=False, default="")
        explanations = db.Column(db.Text, nullable=False, default="{}")
        weights = db.Column(db.Text, nullable=False, default=lambda: json.dumps(DEFAULT_WEIGHTS))
        created_at = db.Column(db.DateTime, nullable=False, default=lambda: datetime.now(timezone.utc))

    class ResumeRewrite(db.Model):
        __tablename__ = "resume_rewrite"
        id = db.Column(db.Integer, primary_key=True)
        user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False, index=True)
        analysis_id = db.Column(db.Integer, db.ForeignKey("analysis.id", ondelete="CASCADE"), nullable=False, index=True)
        tone = db.Column(db.String(30), nullable=False, default="professional")
        original_text = db.Column(db.Text, nullable=False)
        rewritten_text = db.Column(db.Text, nullable=False)
        created_at = db.Column(db.DateTime, nullable=False, default=lambda: datetime.now(timezone.utc))
        updated_at = db.Column(db.DateTime, nullable=False, default=lambda: datetime.now(timezone.utc))

    bp = Blueprint("v2", __name__)

    def get_user_weights(user_id=None):
        if user_id is None:
            if not current_user.is_authenticated:
                return DEFAULT_WEIGHTS.copy()
            user_id = current_user.id
        preference = ScoringPreference.query.filter_by(user_id=user_id).first()
        return _clean_weights(preference.weights if preference else None)

    def weighted_score(score_breakdown, weights=None):
        selected = _clean_weights(weights or get_user_weights())
        score = sum(
            max(0, min(100, int(score_breakdown.get(key, 0)))) * selected[key] / 100
            for key in DEFAULT_WEIGHTS
        )
        return max(0, min(100, round(score)))

    def build_explanations(resume_text, job_description, score_breakdown, matching_skills, missing_skills):
        resume_lines = [line.strip(" -\t") for line in resume_text.splitlines() if line.strip()]
        evidence_lines = []
        for line in resume_lines:
            if any(re.search(rf"(?<!\w){re.escape(skill)}(?!\w)", line, re.I) for skill in matching_skills):
                evidence_lines.append(line[:220])
            if len(evidence_lines) == 4:
                break

        def category(label, score, evidence, missing):
            return {
                "label": label,
                "score": int(score or 0),
                "evidence": evidence[:4] or ["No strong direct evidence was detected in this section."],
                "missing": missing[:4] or ["No major gap detected for this category."],
            }

        exp_section = helpers["extract_experience_section"](resume_text)
        project_section = helpers["extract_projects_section"](resume_text)
        education_section = helpers["extract_education_section"](resume_text)
        return {
            "skills": category(
                "Skills Match", score_breakdown.get("skills"),
                ([f"Matched skill: {skill}" for skill in matching_skills] + evidence_lines),
                [f"Required skill not detected: {skill}" for skill in missing_skills],
            ),
            "experience": category(
                "Experience Relevance", score_breakdown.get("experience"),
                [line.strip()[:220] for line in exp_section.splitlines() if len(line.strip()) > 12],
                ["Add truthful, role-relevant responsibilities.", "Quantify outcomes only when the number is known."],
            ),
            "projects": category(
                "Projects Relevance", score_breakdown.get("projects"),
                [line.strip()[:220] for line in project_section.splitlines() if len(line.strip()) > 12],
                ["State your contribution, technologies and verified outcome for each project."],
            ),
            "education": category(
                "Education Match", score_breakdown.get("education"),
                [line.strip()[:220] for line in education_section.splitlines() if len(line.strip()) > 8],
                ["Include degree, field, institution and graduation year when applicable."],
            ),
        }

    def save_analysis_context(analysis, resume_text, explanations, weights):
        context = AnalysisContext.query.filter_by(analysis_id=analysis.id).first()
        if context is None:
            context = AnalysisContext(user_id=analysis.user_id, analysis_id=analysis.id)
            db.session.add(context)
        context.resume_text = resume_text
        context.explanations = json.dumps(explanations)
        context.weights = json.dumps(_clean_weights(weights))
        return context

    def improve_line(line, tone):
        stripped = line.strip()
        if not stripped:
            return ""
        heading = stripped.rstrip(":").lower()
        known_headings = {
            "summary", "professional summary", "objective", "career objective", "experience",
            "work experience", "employment", "education", "skills", "technical skills",
            "projects", "academic projects", "certifications", "achievements", "languages",
        }
        if heading in known_headings or (len(stripped) < 35 and stripped.isupper()):
            return stripped.rstrip(":").title()

        prefix = ""
        body = stripped
        match = re.match(r"^([•*\-–—]+)\s*(.*)$", stripped)
        if match:
            prefix, body = "• ", match.group(2).strip()

        replacements = [
            (r"^worked on\b", "Contributed to"),
            (r"^responsible for\b", "Managed"),
            (r"^helped (?:to )?\b", "Supported"),
            (r"^made\b", "Developed"),
            (r"^created\b", "Developed"),
            (r"^did\b", "Executed"),
            (r"^used\b", "Applied"),
        ]
        for pattern, replacement in replacements:
            body = re.sub(pattern, replacement, body, count=1, flags=re.I)
        body = re.sub(r"\s+", " ", body).strip()
        if body and body[0].islower():
            body = body[0].upper() + body[1:]
        if tone == "concise":
            body = re.sub(r"\b(very|really|successfully|various|different)\b\s*", "", body, flags=re.I)
        elif tone == "technical" and re.search(r"\b(developed|built|implemented|designed)\b", body, re.I):
            body = re.sub(r"\bproject\b", "solution", body, flags=re.I)
        if body and body[-1] not in ".:;!?":
            body += "."
        return prefix + body

    def rewrite_resume_text(original_text, tone):
        tone = tone if tone in ALLOWED_TONES else "professional"
        improved = [improve_line(line, tone) for line in original_text.splitlines()]
        text = "\n".join(improved)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        return text

    def _analysis_report_data(analysis):
        breakdown = _json(analysis.score_breakdown, {})
        questions = _json(analysis.interview_questions, {})
        matching = _json(analysis.matching_skills, [])
        missing = _json(analysis.missing_skills, [])
        recommendation = helpers["generate_final_recommendation"](analysis.match_score, matching, missing)
        suggestions = helpers["generate_targeted_recommendations"](breakdown, missing)
        suggestions = [
            item.get("message", "") if isinstance(item, dict) else str(item)
            for item in suggestions
        ]
        return {
            "filename": analysis.resume_filename,
            "file_type": "Saved analysis",
            "file_size": "Not stored",
            "word_count": "Not stored",
            "resume_score": analysis.match_score,
            "skills_score": int(breakdown.get("skills", 0)),
            "experience_score": int(breakdown.get("experience", 0)),
            "education_score": int(breakdown.get("education", 0)),
            "projects_score": int(breakdown.get("projects", 0)),
            "ats_label": helpers["calculate_ats_rating"](analysis.match_score).get("label", "Not available"),
            "strength_score": analysis.match_score,
            "strength_label": "Saved analysis",
            "final_title": recommendation.get("title", "Resume Analysis"),
            "final_message": recommendation.get("message", ""),
            "next_action": recommendation.get("next_action", ""),
            "matching_skills": matching,
            "missing_skills": missing,
            "strong_areas": helpers["generate_candidate_strengths"](breakdown, matching),
            "improvement_areas": helpers["generate_candidate_weaknesses"](breakdown, missing),
            "suggestions": [item for item in suggestions if item],
            "technical_questions": questions.get("technical_questions", []),
            "resume_questions": questions.get("resume_questions", []),
            "job_questions": questions.get("job_questions", []),
        }

    @bp.route("/reports")
    @login_required
    def reports():
        query = Analysis.query.filter_by(user_id=current_user.id)
        search = request.args.get("q", "").strip()
        analysis_type = request.args.get("type", "all")
        recommendation = request.args.get("recommendation", "all")
        try:
            minimum_score = max(0, min(100, int(request.args.get("min_score", 0))))
        except ValueError:
            minimum_score = 0
        if search:
            query = query.filter(Analysis.resume_filename.ilike(f"%{search}%"))
        if analysis_type in {"single", "multiple"}:
            query = query.filter_by(analysis_type=analysis_type)
        query = query.filter(Analysis.match_score >= minimum_score)
        sort = request.args.get("sort", "newest")
        order = Analysis.match_score.desc() if sort == "score" else Analysis.created_at.desc()
        analyses = query.order_by(order).all()
        if recommendation in {"strong", "potential", "low"}:
            ranges = {"strong": (75, 100), "potential": (50, 74), "low": (0, 49)}
            low, high = ranges[recommendation]
            analyses = [item for item in analyses if low <= item.match_score <= high]
        return render_template(
            "reports.html", analyses=analyses, search=search, analysis_type=analysis_type,
            minimum_score=minimum_score, recommendation=recommendation, sort=sort,
        )

    @bp.route("/reports/<int:analysis_id>/download")
    @login_required
    def download_saved_report(analysis_id):
        analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first_or_404()
        buffer = helpers["build_analysis_pdf"](_analysis_report_data(analysis))
        safe_name = secure_filename(analysis.resume_filename.rsplit(".", 1)[0]) or "resume"
        return send_file(buffer, mimetype="application/pdf", as_attachment=True, download_name=f"{safe_name}_saved_report.pdf")

    @bp.route("/reports/<int:analysis_id>/delete", methods=["POST"])
    @login_required
    def delete_saved_report(analysis_id):
        analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first_or_404()
        AnalysisContext.query.filter_by(analysis_id=analysis.id, user_id=current_user.id).delete()
        ResumeRewrite.query.filter_by(analysis_id=analysis.id, user_id=current_user.id).delete()
        db.session.delete(analysis)
        db.session.commit()
        flash("Saved report deleted successfully.", "success")
        return redirect(url_for("v2.reports"))

    @bp.route("/scoring-weights", methods=["GET", "POST"])
    @login_required
    def scoring_weights():
        weights = get_user_weights()
        if request.method == "POST":
            try:
                submitted = {key: int(request.form.get(key, "0")) for key in DEFAULT_WEIGHTS}
            except ValueError:
                submitted = {}
            if any(value < 0 or value > 100 for value in submitted.values()) or sum(submitted.values()) != 100:
                flash("Each weight must be between 0 and 100, and the total must equal 100%.", "danger")
                return render_template("scoring_weights.html", weights=submitted or weights)
            preference = ScoringPreference.query.filter_by(user_id=current_user.id).first()
            if preference is None:
                preference = ScoringPreference(user_id=current_user.id)
                db.session.add(preference)
            preference.weights = json.dumps(submitted)
            preference.updated_at = datetime.now(timezone.utc)
            db.session.commit()
            flash("Scoring weights saved. They now apply to single and multiple-resume analysis.", "success")
            return redirect(url_for("v2.scoring_weights"))
        return render_template("scoring_weights.html", weights=weights)

    @bp.route("/rewrite/<int:analysis_id>", methods=["GET", "POST"])
    @login_required
    def rewrite_resume(analysis_id):
        analysis = Analysis.query.filter_by(id=analysis_id, user_id=current_user.id).first_or_404()
        context = AnalysisContext.query.filter_by(analysis_id=analysis.id, user_id=current_user.id).first()
        if context is None or not context.resume_text.strip():
            flash("Resume rewriting is available for analyses created after the Version 2 upgrade. Analyze the resume again first.", "warning")
            return redirect(url_for("analysis_details", analysis_id=analysis.id))
        rewrite = ResumeRewrite.query.filter_by(analysis_id=analysis.id, user_id=current_user.id).order_by(ResumeRewrite.updated_at.desc()).first()
        if request.method == "POST":
            action = request.form.get("action", "generate")
            tone = request.form.get("tone", "professional")
            if tone not in ALLOWED_TONES:
                tone = "professional"
            if action == "save":
                rewritten = request.form.get("rewritten_text", "").strip()
                if len(rewritten.split()) < 3:
                    flash("The rewritten resume cannot be empty.", "danger")
                    return redirect(url_for("v2.rewrite_resume", analysis_id=analysis.id))
            else:
                rewritten = rewrite_resume_text(context.resume_text, tone)
            if rewrite is None:
                rewrite = ResumeRewrite(
                    user_id=current_user.id, analysis_id=analysis.id,
                    original_text=context.resume_text, rewritten_text=rewritten, tone=tone,
                )
                db.session.add(rewrite)
            else:
                rewrite.rewritten_text = rewritten
                rewrite.tone = tone
                rewrite.updated_at = datetime.now(timezone.utc)
            db.session.commit()
            flash("Resume rewrite saved. Review every statement before downloading.", "success")
            return redirect(url_for("v2.rewrite_resume", analysis_id=analysis.id))
        explanations = _json(context.explanations, {})
        return render_template("rewrite.html", analysis=analysis, context=context, rewrite=rewrite, explanations=explanations)

    @bp.route("/rewrite/<int:rewrite_id>/download/<file_format>")
    @login_required
    def download_rewrite(rewrite_id, file_format):
        rewrite = ResumeRewrite.query.filter_by(id=rewrite_id, user_id=current_user.id).first_or_404()
        analysis = Analysis.query.filter_by(id=rewrite.analysis_id, user_id=current_user.id).first_or_404()
        safe_name = secure_filename(analysis.resume_filename.rsplit(".", 1)[0]) or "resume"
        if file_format == "docx":
            document = Document()
            for line in rewrite.rewritten_text.splitlines():
                clean = line.strip()
                if not clean:
                    document.add_paragraph("")
                elif clean.startswith("• "):
                    document.add_paragraph(clean[2:], style="List Bullet")
                elif len(clean) < 40 and clean.title() == clean:
                    document.add_heading(clean, level=2)
                else:
                    document.add_paragraph(clean)
            buffer = BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return send_file(buffer, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name=f"{safe_name}_rewritten.docx")
        if file_format == "pdf":
            buffer = BytesIO()
            styles = getSampleStyleSheet()
            story = [Paragraph("Improved Resume", styles["Title"]), Spacer(1, 12)]
            for line in rewrite.rewritten_text.splitlines():
                clean = line.strip()
                if clean:
                    story.append(Paragraph(clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["BodyText"]))
                    story.append(Spacer(1, 5))
            SimpleDocTemplate(buffer, pagesize=A4, title="Improved Resume").build(story)
            buffer.seek(0)
            return send_file(buffer, mimetype="application/pdf", as_attachment=True, download_name=f"{safe_name}_rewritten.pdf")
        flash("Unsupported download format.", "danger")
        return redirect(url_for("v2.rewrite_resume", analysis_id=rewrite.analysis_id))

    @bp.route("/rewrite/<int:rewrite_id>/reanalyze", methods=["POST"])
    @login_required
    def reanalyze_rewrite(rewrite_id):
        rewrite = ResumeRewrite.query.filter_by(id=rewrite_id, user_id=current_user.id).first_or_404()
        source = Analysis.query.filter_by(id=rewrite.analysis_id, user_id=current_user.id).first_or_404()
        text = rewrite.rewritten_text
        job_description = source.job_description
        resume_skills = helpers["extract_skills"](text)
        job_skills = helpers["extract_skills"](job_description)
        matching = sorted(set(resume_skills) & set(job_skills))
        missing = sorted(set(job_skills) - set(resume_skills))
        breakdown = helpers["calculate_score_breakdown"](text, job_description, matching, job_skills)
        weights = get_user_weights()
        score = weighted_score(breakdown, weights)
        questions = helpers["generate_interview_questions"](text, job_description, resume_skills, matching, missing, job_skills)
        analysis = Analysis(
            user_id=current_user.id, resume_filename=f"rewritten_{source.resume_filename}",
            analysis_type="single", job_description=job_description, match_score=score,
            detected_skills=json.dumps(resume_skills), matching_skills=json.dumps(matching),
            missing_skills=json.dumps(missing), candidate_rank=None,
            interview_questions=json.dumps(questions), score_breakdown=json.dumps(breakdown),
        )
        db.session.add(analysis)
        db.session.flush()
        explanations = build_explanations(text, job_description, breakdown, matching, missing)
        save_analysis_context(analysis, text, explanations, weights)
        db.session.commit()
        return render_template(
            "reanalyze_result.html", source=source, analysis=analysis, score_breakdown=breakdown,
            matching_skills=matching, missing_skills=missing, explanations=explanations, weights=weights,
        )

    app.register_blueprint(bp)

    return SimpleNamespace(
        get_user_weights=get_user_weights,
        weighted_score=weighted_score,
        build_explanations=build_explanations,
        save_analysis_context=save_analysis_context,
        models={
            "ScoringPreference": ScoringPreference,
            "AnalysisContext": AnalysisContext,
            "ResumeRewrite": ResumeRewrite,
        },
    )
